from pdfminer.high_level import extract_text
from docx import Document
import tempfile
from io import BytesIO
import re

class FileHandler:
    @staticmethod
    def read_file(uploaded_file):
        file_ext = uploaded_file.name.split('.')[-1].lower()
        content = uploaded_file.getvalue()
        
        if file_ext == 'pdf':
            return FileHandler._read_pdf(content)
        elif file_ext == 'docx':
            return FileHandler._read_docx(content)
        elif file_ext == 'txt':
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                return content.decode('latin-1')
        else:
            raise ValueError(f"Formato no soportado: {file_ext}")
    
    @staticmethod
    def _read_pdf(content):
        with tempfile.NamedTemporaryFile(delete=True, suffix='.pdf') as tmp:
            tmp.write(content)
            text = extract_text(tmp.name)
            return text.replace('\x0c', '\n\n[PAGE_BREAK]\n\n')
    
    @staticmethod
    def _read_docx(content):
        with BytesIO(content) as bio:
            doc = Document(bio)
            return '\n\n'.join([para.text for para in doc.paragraphs])
    
    @staticmethod
    def create_output_file(content, original_name, target_language):
        # El nombre del archivo ahora incluye el idioma de destino
        output_name = f"TRADUCCION_LITERARIA_{target_language}_{original_name}"
        doc = Document()
        
        for para in content.split('\n\n'):
            doc.add_paragraph(para)
        
        bio = BytesIO()
        doc.save(bio)
        
        if not output_name.endswith('.docx'):
            output_name = output_name.rsplit('.', 1)[0] + '.docx'
        
        return bio.getvalue(), output_name
